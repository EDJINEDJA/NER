import re
import os
import logging
import pdfplumber
from striprtf.striprtf import rtf_to_text
from docx import Document

class PdfDocxDoc2Txt():
    
    # Convert .rtf File To Text
    def convert_rtf2text(self,word_file):
        """
            A utility function to convert a Microsoft rtf files to raw text.
            This code is largely borrowed from existing solutions, and does not match the style of the rest of this repo.
            :param rtf_file: rtf file with gets uploaded by the user
            :type rtf_file: InMemoryUploadedFile
            :return: The text contents of the rtf file
            :rtype: str
        """
 
        pattern0="\[\d+\:\d+\:\d+\.\d+\]|\[\d+]|\--|Scribe3|Date: \d+\-\d+\-\d+|Document: \d+\-\d+.wma|Duration: \d+.\d+s|Language: fre \(auto\)|Type: cts"
        pattern1="\\n"
        pattern2="           "
        pattern3='-'
        with open(word_file,encoding="UTF-8") as infile:
            content = infile.read()
            text = rtf_to_text(content)
            text = re.sub(pattern0,'',text)
            text = re.sub(pattern1,'   ',text)
            text = re.sub(pattern2,'.',text)
            text = re.sub(pattern3 ,'',text)
        return text
    
    # Convert .docx File To Text
    def convert_docx2txt(self, docx_file):
        """
            A utility function to convert a Microsoft docx files to raw text.
            This code is largely borrowed from existing solutions, and does not match the style of the rest of this repo.
            :param docx_file: docx file with gets uploaded by the user
            :type docx_file: InMemoryUploadedFile
            :return: The text contents of the docx file
            :rtype: str
        """
        
        pattern0="\[\d+\:\d+\:\d+\.\d+\]|\[\d+]|\--|Scribe3|Date: \d+\-\d+\-\d+|Document: \d+\-\d+.wma|Duration: \d+.\d+s|Language: fre \(auto\)|Type: cts"
        pattern1="\\n"
        pattern2="           "
        pattern3='-'
        
        doc = Document(docx_file)
        content = ' '.join([paragraph.text for paragraph in doc.paragraphs])

        # replace following line with location of your .docx file
        content = doc.get_text()
        
        try:
            # text = rtf_to_text(content)
            text = re.sub(pattern0,'',content)
            text = re.sub(pattern1,'   ',text)
            text = re.sub(pattern2,'.',text)
            text = re.sub(pattern3 ,'',text)
            clean_text = re.sub(r'\n+', '\n', text)
            clean_text = clean_text.replace("\r", "\n").replace("\t", " ")  # Normalize text blob
            resume_lines = clean_text.splitlines()  # Split text blob into individual lines
            resume_lines = [re.sub('\s+', ' ', line.strip()) for line in resume_lines if
                            line.strip()]  # Remove empty strings and whitespaces
            return resume_lines, text
        except Exception as e:
            logging.error('Error in docx file:: ' + str(e))
            return [], " "

    # Convert .doc File To Text
    def convert_doc2txt(self, docx_file):
        """
            A utility function to convert a Microsoft doc files to raw text.
            This code is largely borrowed from existing solutions, and does not match the style of the rest of this repo.
            :param doc_file: doc file with gets uploaded by the user
            :type doc_file: InMemoryUploadedFile
            :return: The text contents of the doc file
            :rtype: str
        """
        
        pattern0="\[\d+\:\d+\:\d+\.\d+\]|\[\d+]|\--|Scribe3|Date: \d+\-\d+\-\d+|Document: \d+\-\d+.wma|Duration: \d+.\d+s|Language: fre \(auto\)|Type: cts"
        pattern1="\\n"
        pattern2="           "
        pattern3='-'
        
        # replace following line with location of your .doc file
        doc = Document(docx_file)
        content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
        
        try:
            # text = rtf_to_text(content)
            text = re.sub(pattern0,'',content)
            text = re.sub(pattern1,'   ',text)
            text = re.sub(pattern2,'.',text)
            text = re.sub(pattern3 ,'',text)
            clean_text = re.sub(r'\n+', '\n', text)
            clean_text = clean_text.replace("\r", "\n").replace("\t", " ")  # Normalize text blob
            resume_lines = clean_text.splitlines()  # Split text blob into individual lines
            resume_lines = [re.sub('\s+', ' ', line.strip()) for line in resume_lines if
                            line.strip()]  # Remove empty strings and whitespaces
            return resume_lines, text
        except Exception as e:
            logging.error('Error in docx file:: ' + str(e))
            return [], " "
        
    # Convert .pdf File To Text
    def convert_pdf2txt(self, pdf_file):
        """
        A utility function to convert a machine-readable PDF to raw text.
        This code is largely borrowed from existing solutions, and does not match the style of the rest of this repo.
        :param input_pdf_path: Path to the .pdf file which should be converted
        :type input_pdf_path: str
        :return: The text contents of the pdf
        :rtype: str
        """

        pdf = pdfplumber.open(pdf_file)
        raw_text= ""
       
        for page in pdf.pages:
            raw_text += page.extract_text() + "\n"

        pdf.close()                
      
        try:
            full_string = re.sub(r'\n+', '\n', raw_text)
            full_string = full_string.replace("\r", "\n")
            full_string = full_string.replace("\t", " ")

            # Remove awkward LaTeX bullet characters
            full_string = re.sub(r"\uf0b7", " ", full_string)
            full_string = re.sub(r"\(cid:\d{0,3}\)", " ", full_string)
            full_string = re.sub(r'• ', " ", full_string)

            # Split text blob into individual lines
            resume_lines = full_string.splitlines(True)

            # Remove empty strings and whitespaces
            resume_lines = [re.sub('\s+', ' ', line.strip()) for line in resume_lines if line.strip()]
           
            return resume_lines, raw_text 
        except Exception as e:
            logging.error('Error in docx file:: ' + str(e))
            return [], " "

